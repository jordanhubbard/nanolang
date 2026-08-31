#!/usr/bin/env python3
"""Build the NanoLang developer narrative from repository evidence."""

from __future__ import annotations

import json
import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt, RGBColor

HERE = Path(__file__).resolve().parent
REPO = Path(os.environ.get("NANOLANG_DECK_REPO", str(HERE.parents[1])))
OUT = Path(os.environ.get("NANOLANG_NARRATIVE_OUTPUT", str(HERE / "nanolang-developer-overview.docx")))
PPTX = Path(os.environ.get("NANOLANG_DECK_OUTPUT", str(HERE / "nanolang-developer-overview.pptx")))


def paragraph(document: Document, value: str, *, style: str | None = None, code: bool = False) -> None:
    item = document.add_paragraph(style=style)
    item.paragraph_format.space_after = Pt(8)
    item.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = item.add_run(value)
    run.font.name = "Courier New" if code else "Calibri"
    run.font.size = Pt(9 if code else 11)
    run.font.color.rgb = RGBColor(0x65, 0x70, 0x7C) if code else RGBColor(0x10, 0x13, 0x17)


def heading(document: Document, level: int, value: str) -> None:
    document.add_heading(value, level=level)


def build() -> Path:
    document = Document()
    section = document.sections[0]
    section.left_margin = section.right_margin = Inches(0.9)
    document.core_properties.title = "NanoLang: the language, compiler, and VM"
    document.core_properties.author = "NanoLang"
    document.core_properties.subject = "NanoLang 3.5 developer narrative"

    heading(document, 1, "NanoLang: the language, compiler, and VM")
    paragraph(document, "I am NanoLang. This is my developer narrative for release 3.5. It explains my language contract, compiler paths, NanoISA bytecode, NanoVM execution, foreign-function boundary, tests, diagnostics, and the work I have reserved for 4.0.")
    paragraph(document, "Authority: docs/PERSONA.md, README.md, docs/NANOISA.md, docs/ROADMAP.md, spec/nanoisa.yaml, and the current test suites.")

    heading(document, 1, "The language contract")
    heading(document, 2, "Explicit syntax and types")
    paragraph(document, "My function calls use prefix form. My function boundaries use explicit parameter and result types. My operators do not rely on a hidden precedence table. These choices make source easier to parse and review.")
    paragraph(document, "fn gcd(a: int, b: int) -> int {\n    return a\n}\n\nshadow gcd {\n    assert (== (gcd 48 18) 6)\n}", code=True)
    heading(document, 2, "Shadow tests")
    paragraph(document, "Every eligible function carries a shadow test. The test is an executable statement about behavior. It is not a coverage decoration and it is not a substitute for integration tests.")
    heading(document, 2, "The formal core and its boundary")
    paragraph(document, "My NanoCore semantics have mechanized proofs for preservation, progress, determinism, semantic equivalence, and evaluator soundness. Full NanoLang, NanoISA, FFI, and generated-C behavior remains outside that proved subset unless explicitly stated otherwise.")

    heading(document, 1, "The compilation paths")
    heading(document, 2, "C transpilation")
    paragraph(document, "nanoc lowers NanoLang to generated C and links the runtime libraries. This is my native production path. The generated program carries the behavior of the source only after compilation and its tests pass.")
    heading(document, 2, "NanoISA generation")
    paragraph(document, "nano_virt lowers NanoLang to a serialized .nvm module. The active opcode metadata comes from spec/nanoisa.yaml and is generated into src/nanoisa/generated_schema.h.")
    paragraph(document, ".nano source  ->  nano_virt  ->  .nvm module  ->  nano_vm", code=True)
    heading(document, 2, "NanoVM execution")
    paragraph(document, "NanoVM decodes each function once, records instruction boundaries, and executes the decoded representation. A private indexed cursor reduces repeated byte-offset lookup while preserving byte offsets for diagnostics and traps.")
    heading(document, 2, "Side-table debug data")
    paragraph(document, "NanoVirt records source locations in side tables. It does not emit executable debug-line instructions into normal output. --strip-debug removes those side tables when they are not wanted.")

    heading(document, 1, "Runtime boundaries")
    heading(document, 2, "Heap values and ownership")
    paragraph(document, "NanoVM values are tagged. Strings, arrays, structs, tuples, unions, hash maps, and closures are heap objects managed through retain and release operations. The remaining 4.0 roadmap includes additional ownership and representation work; I do not describe that work as finished.")
    heading(document, 2, "FFI and the co-process boundary")
    paragraph(document, "Extern declarations become typed import records. NanoVM routes calls through traps and can isolate foreign calls in nano_cop. The co-process wire fields are explicitly little-endian in the current implementation.")
    heading(document, 2, "Modules and serialized bytecode")
    paragraph(document, "The current NVM loader rejects duplicate or overlapping sections, directory intrusion, partial fixed-width records, trailing data, and arithmetic-overflow ranges. The broader v2 module contract remains 4.0 work.")

    heading(document, 1, "Evidence and diagnostics")
    heading(document, 2, "Tests")
    paragraph(document, "The current focused suites report 1,090 NanoISA tests, 322 NanoVM tests, and 63 NanoVirt tests passing. CI additionally exercises x64, arm64, C, PTX, RISC-V, sanitizers, coverage, documentation, benchmarks, and security checks.")
    heading(document, 2, "NanoISA profiles and opcode traces")
    paragraph(document, "--profile-isa writes structured counters for retired instructions, opcode sequences, branches, calls, stack and frame depth, traps, heap traffic, and FFI traffic. NANO_VM_TRACE is read once during VM initialization and enables per-instruction records with opcode, function, offset, stack values, and FFI results.")
    heading(document, 2, "Generated-C profiling")
    paragraph(document, "Generated timing hooks read NANO_PROFILE once when the executable starts. NANO_PROFILE=0 disables a --profile build without rebuilding it. Disabled hooks do not perform environment reads or timing work at each event.")

    heading(document, 1, "Release 3.5")
    heading(document, 2, "What I shipped")
    paragraph(document, "I shipped measured NanoISA execution, generated ISA metadata, typed scalar and aggregate operations, side-table debug data, predecoded execution, direct tail calls, opcode diagnostics, and runtime-selectable generated-C profiling.")
    heading(document, 2, "What I measured")
    paragraph(document, "The benchmark harness runs seven maintained workloads with 20 samples per workload. All seven workloads completed with exit code zero. The summary records distributions and normalized cost rather than a single timing claim.")
    heading(document, 2, "What I do not claim")
    paragraph(document, "I do not claim that my full NanoISA v2 verifier, linked callable handles, v2 module header, typed mixed ABI, runtime representation redesign, computed-goto dispatch, or superinstruction program is complete. Those are 4.0 roadmap work.")

    heading(document, 1, "The 4.0 boundary")
    heading(document, 2, "Verifier and module work")
    paragraph(document, "I still need stack-height and type-flow verification across control-flow joins, complete signature and ownership checks, richer module sections, symbolic assembly operands, lossless disassembly, and exhaustive opcode coverage.")
    heading(document, 2, "Runtime representation and dispatch")
    paragraph(document, "I still need measured decisions for globals, constants, strings, homogeneous arrays, hash maps, ownership, and private optimized dispatch. A faster path earns its place only through repeatable NanoLang and Forth measurements.")
    heading(document, 2, "FFI and transport")
    paragraph(document, "I still need typed call descriptors, a sound mixed-type ABI, consistent argument limits, large-payload transport, and lifecycle evidence for crashes, restarts, and batching.")

    heading(document, 1, "How to work on me")
    heading(document, 2, "Read the source and roadmap")
    paragraph(document, "Start with docs/PERSONA.md, docs/ROADMAP.md, the relevant source symbols, and the matching tests. Keep 4.0 work in its declared queue. Do not turn a roadmap sentence into a feature claim.")
    heading(document, 2, "Run the gates")
    paragraph(document, "make -f Makefile.gnu test-nanoisa\nmake -f Makefile.gnu test-nanovm\nmake -f Makefile.gnu test-nanovirt\nmake -f Makefile.gnu test-dynamic-profile\nNANOISA_BENCH_RUNS=20 bash scripts/benchmark_nanoisa.sh", code=True)
    paragraph(document, "I say what I mean, I show what I tested, and I leave the unproved boundary visible.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUT)
    return OUT


def main() -> None:
    output = build()
    if not PPTX.is_file():
        raise SystemExit(f"presentation artifact missing: {PPTX}")
    manifest = Path(os.environ.get("OBJ_DIR", str(REPO / "_build"))) / "nanolang-developer-overview" / "capability-manifest.json"
    manifest.parent.mkdir(parents=True, exist_ok=True)
    manifest.write_text(json.dumps({"schema": "nanolang/developer-document-pair@1", "slides": 12, "narrative": str(output)}, indent=2) + "\n")
    print(f"built narrative -> {output}")


if __name__ == "__main__":
    main()
